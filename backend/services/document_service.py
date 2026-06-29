import io
import mimetypes
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from bson import ObjectId
from docx import Document
from flask import current_app
from gridfs import GridFSBucket
from pymongo import MongoClient
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename


TEMPLATE_DIRECTORY = Path(__file__).resolve().parent.parent / "document_templates"
TEMPLATES = {
    "proposal": {"name": "Project Proposal", "filename": "proposal.docx"},
    "terms": {"name": "Terms of Contract", "filename": "terms-of-contract.docx"},
    "maintenance": {"name": "Support and Maintenance Agreement", "filename": "maintenance-agreement.docx"},
    "handover": {"name": "Project Handover", "filename": "handover.docx"},
}
ALLOWED_EXTENSIONS = {".doc", ".docx", ".pdf", ".xls", ".xlsx", ".csv", ".txt", ".png", ".jpg", ".jpeg"}
MAX_FILE_BYTES = 8 * 1024 * 1024
_client = None


class DocumentStorageError(RuntimeError):
    pass


def list_document_templates() -> list[dict[str, str]]:
    return [{"id": key, "name": value["name"], "filename": value["filename"]} for key, value in TEMPLATES.items()]


def list_documents(owner_type: str, owner_id: str) -> list[dict[str, Any]]:
    query = {"owner_type": owner_type, "owner_id": owner_id}
    return _serialise_many(_metadata().find(query, {"_id": 0, "gridfs_id": 0}).sort("created_at", -1))


def list_correspondence_documents(enquiry: dict[str, Any]) -> list[dict[str, Any]]:
    email = str(enquiry.get("email") or "").lower()
    query = {"$or": [{"owner_type": "enquiry", "owner_id": enquiry["id"]}, {"customer_email": email}]}
    return _serialise_many(_metadata().find(query, {"_id": 0, "gridfs_id": 0}).sort("created_at", -1))


def generate_document(template_id: str, owner_type: str, owner_id: str) -> dict[str, Any]:
    template = TEMPLATES.get(template_id)
    if not template:
        raise ValueError("Unknown document template.")
    project, customer = _owner_context(owner_type, owner_id)
    document = Document(TEMPLATE_DIRECTORY / template["filename"])
    _populate_document(document, template_id, project, customer)
    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    project_name = str(project.get("name") or customer.get("name") or "client")
    filename = secure_filename(f"{project_name}-{template_id}.docx")
    return _store_bytes(stream.read(), filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", owner_type, owner_id, str(project.get("client_email") or customer.get("email") or ""), template_id)


def store_uploaded_document(file: FileStorage, owner_type: str, owner_id: str, customer_email: str = "") -> dict[str, Any]:
    if owner_type in {"project", "customer"}:
        project, customer = _owner_context(owner_type, owner_id)
        customer_email = str(project.get("client_email") or customer.get("email") or "")
    elif owner_type != "enquiry":
        raise ValueError("Invalid document owner.")
    filename = secure_filename(file.filename or "")
    extension = Path(filename).suffix.lower()
    if not filename or extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported attachment type.")
    content = file.read(MAX_FILE_BYTES + 1)
    if not content or len(content) > MAX_FILE_BYTES:
        raise ValueError("Attachments must contain data and be no larger than 8 MB.")
    content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    return _store_bytes(content, filename, content_type, owner_type, owner_id, customer_email)


def store_invoice_pdf(project: dict[str, Any], invoice: dict[str, Any], content: bytes) -> dict[str, Any]:
    project_id = str(project.get("id") or "").strip()
    invoice_id = str(invoice.get("id") or "").strip()
    if not project_id or not invoice_id:
        raise ValueError("A saved project and invoice are required.")
    template_id = f"invoice:{invoice_id}"
    existing = _metadata().find_one(
        {"owner_type": "project", "owner_id": project_id, "template_id": template_id},
        {"id": 1},
    )
    if existing:
        delete_document(str(existing["id"]))
    reference = secure_filename(str(invoice.get("reference") or "invoice"))[:80] or "invoice"
    return _store_bytes(
        content,
        f"invoice-{reference}.pdf",
        "application/pdf",
        "project",
        project_id,
        str(project.get("client_email") or ""),
        template_id,
    )


def get_document(document_id: str) -> tuple[dict[str, Any], bytes] | None:
    record = _metadata().find_one({"id": document_id}, {"_id": 0})
    if not record:
        return None
    stream = io.BytesIO()
    _bucket().download_to_stream(ObjectId(record["gridfs_id"]), stream)
    record.pop("gridfs_id", None)
    return _serialise(record), stream.getvalue()


def delete_document(document_id: str) -> bool:
    record = _metadata().find_one({"id": document_id})
    if not record:
        return False
    _bucket().delete(ObjectId(record["gridfs_id"]))
    _metadata().delete_one({"id": document_id})
    return True


def delete_customer_documents(email: str) -> int:
    clean_email = str(email or "").strip().lower()
    records = list(_metadata().find({"customer_email": clean_email}, {"id": 1}))
    deleted = 0
    for record in records:
        if delete_document(str(record.get("id") or "")):
            deleted += 1
    return deleted


def delete_owner_documents(owner_type: str, owner_id: str) -> int:
    records = list(_metadata().find({"owner_type": owner_type, "owner_id": owner_id}, {"id": 1}))
    deleted = 0
    for record in records:
        if delete_document(str(record.get("id") or "")):
            deleted += 1
    return deleted


def load_email_attachments(enquiry: dict[str, Any], document_ids: list[str]) -> list[dict[str, Any]]:
    allowed = {item["id"] for item in list_correspondence_documents(enquiry)}
    attachments = []
    for document_id in document_ids[:10]:
        if document_id not in allowed:
            raise ValueError("An attachment is not available for this customer.")
        loaded = get_document(document_id)
        if loaded:
            metadata, content = loaded
            attachments.append({**metadata, "content": content})
    if sum(len(item["content"]) for item in attachments) > MAX_FILE_BYTES:
        raise ValueError("Combined attachments must be no larger than 8 MB.")
    return attachments


def _store_bytes(content: bytes, filename: str, content_type: str, owner_type: str, owner_id: str, customer_email: str, template_id: str = "") -> dict[str, Any]:
    file_id = _bucket().upload_from_stream(filename, content, metadata={"content_type": content_type})
    record = {
        "id": str(uuid.uuid4()), "owner_type": owner_type, "owner_id": owner_id,
        "customer_email": customer_email.lower(), "template_id": template_id,
        "kind": "generated" if template_id else "uploaded", "filename": filename,
        "content_type": content_type, "size": len(content), "gridfs_id": str(file_id),
        "created_at": datetime.now(timezone.utc),
    }
    _metadata().insert_one(record)
    record.pop("_id", None)
    record.pop("gridfs_id", None)
    return _serialise(record)


def _owner_context(owner_type: str, owner_id: str) -> tuple[dict[str, Any], dict[str, Any]]:
    database = _database()
    if owner_type == "project":
        project = database[current_app.config["MONGODB_PROJECT_COLLECTION"]].find_one({"id": owner_id}, {"_id": 0})
        if not project:
            raise ValueError("Project not found.")
        email = str(project.get("client_email") or "").lower()
        customer = database[current_app.config["MONGODB_CUSTOMER_COLLECTION"]].find_one({"email": email}, {"_id": 0}) or {"email": email, "name": project.get("client_name", "")}
        return project, customer
    if owner_type == "customer":
        email = owner_id.lower()
        customer = database[current_app.config["MONGODB_CUSTOMER_COLLECTION"]].find_one({"email": email}, {"_id": 0})
        if not customer:
            enquiry = database[current_app.config["MONGODB_ENQUIRY_COLLECTION"]].find_one({"email": email}, {"_id": 0}, sort=[("created_at", -1)]) or {}
            customer = {"email": email, "name": enquiry.get("name") or email}
        return {}, customer
    raise ValueError("Invalid document owner.")


def _populate_document(document: Document, template_id: str, project: dict[str, Any], customer: dict[str, Any]) -> None:
    project_name = str(project.get("name") or "")
    client_name = str(project.get("client_name") or customer.get("name") or "")
    company = str(customer.get("organisation") or "")
    today = datetime.now().strftime("%d/%m/%Y")
    _set_below(document, "Project name:", project_name)
    _set_below(document, "Client name:", client_name)
    _set_below(document, "Client:", client_name)
    _set_below(document, "Handover Date:", "")
    _set_below(document, "Effective Date:", "")
    _set_next(document, "Client:", client_name)
    _set_next(document, "Project:", project_name)
    _set_next(document, "Client Name:", client_name)
    _set_next(document, "Company (if applicable):", company)
    _set_next(document, "Name", client_name, table_hint="Client")
    _set_next(document, "Company", company, table_hint="Client")

    if template_id == "proposal":
        _populate_proposal(document, project, customer, today)
    elif template_id == "handover":
        _populate_handover(document, project, client_name)
    elif template_id == "terms":
        _clear_signature_fields(document)
    elif template_id == "maintenance":
        _clear_signature_fields(document)


def _populate_proposal(document: Document, project: dict[str, Any], customer: dict[str, Any], today: str) -> None:
    values = {"BUSINESS NAME": customer.get("organisation", ""), "PRIMARY CONTACT": project.get("client_name") or customer.get("name", ""), "TELEPHONE NUMBER": customer.get("phone", ""), "EMAIL ADDRESS": project.get("client_email") or customer.get("email", "")}
    for label, value in values.items():
        _set_below(document, label, str(value or ""))
    services = project.get("services") or []
    if len(document.tables) > 9:
        table = document.tables[9]
        for index, service in enumerate(services[:16], start=2):
            if index < len(table.rows):
                _set_cell(table, index, 2, str(service.get("category") or ""))
                _set_cell(table, index, 3, str(service.get("service") or ""))
    if len(document.tables) > 30:
        table = document.tables[30]
        for index, service in enumerate(services[:10], start=1):
            if index < len(table.rows):
                _set_cell(table, index, 1, str(service.get("category") or ""))
                _set_cell(table, index, 3, str(service.get("service") or ""))
    _set_document_control_date(document, today)
    _clear_signature_fields(document)


def _populate_handover(document: Document, project: dict[str, Any], client_name: str) -> None:
    if len(document.tables) > 2 and len(document.tables[2].rows) > 1:
        row = document.tables[2].rows[1]
        for index, value in enumerate([project.get("name", ""), client_name, "", project.get("id", "")]):
            _set_cell(document.tables[2], 1, index, str(value or ""))
    if len(document.tables) > 4:
        for index, service in enumerate((project.get("services") or [])[:5], start=1):
            _set_cell(document.tables[4], index, 1, str(service.get("service") or ""))
            _set_cell(document.tables[4], index, 2, "")
    _clear_signature_fields(document)


def _set_below(document: Document, label: str, value: str) -> None:
    for table in document.tables:
        for row_index, row in enumerate(table.rows[:-1]):
            for column, cell in enumerate(row.cells):
                if cell.text.strip().casefold() == label.casefold() and column < len(table.rows[row_index + 1].cells):
                    _set_cell(table, row_index + 1, column, value)
                    return


def _set_next(document: Document, label: str, value: str, table_hint: str = "") -> None:
    for table in document.tables:
        if table_hint and table_hint.casefold() not in " ".join(cell.text for row in table.rows[:1] for cell in row.cells).casefold():
            continue
        for row_index, row in enumerate(table.rows):
            for column, cell in enumerate(row.cells[:-1]):
                if cell.text.strip().casefold() == label.casefold():
                    _set_cell(table, row_index, column + 1, value)
                    return


def _clear_signature_fields(document: Document) -> None:
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            if row.cells and row.cells[0].text.strip().casefold() in {"signature", "signature:", "date", "date:"} and len(row.cells) > 1:
                _set_cell(table, row_index, 1, "")


def _set_document_control_date(document: Document, value: str) -> None:
    for table in document.tables:
        if table.rows and "Document Control" in " ".join(cell.text for cell in table.rows[0].cells) and len(table.rows) > 2:
            _set_cell(table, 2, 1, value)


def _set_cell(table, row: int, column: int, value: str) -> None:
    if row < len(table.rows) and column < len(table.rows[row].cells):
        table.rows[row].cells[column].text = value


def _database():
    global _client
    if _client is None:
        _client = MongoClient(current_app.config["MONGODB_URI"], serverSelectionTimeoutMS=current_app.config["MONGODB_SERVER_SELECTION_TIMEOUT_MS"])
    return _client[current_app.config["MONGODB_DATABASE"]]


def _metadata():
    return _database()[current_app.config["MONGODB_DOCUMENT_COLLECTION"]]


def _bucket():
    return GridFSBucket(_database(), bucket_name=current_app.config["MONGODB_FILE_BUCKET"])


def _serialise(record: dict[str, Any]) -> dict[str, Any]:
    return {key: value.isoformat().replace("+00:00", "Z") if isinstance(value, datetime) else value for key, value in record.items()}


def _serialise_many(records) -> list[dict[str, Any]]:
    return [_serialise(record) for record in records]
