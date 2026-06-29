import unittest

from docx import Document

from backend.services.document_service import TEMPLATE_DIRECTORY, TEMPLATES, _populate_document


class DocumentTemplatePopulationTests(unittest.TestCase):
    def setUp(self):
        self.project = {
            "id": "PROJ-1", "name": "Example Website", "client_name": "Alex Morgan",
            "client_email": "alex@example.com", "stage": "active",
            "services": [
                {"service": "Website design", "category": "Digital Presence"},
                {"service": "CMS setup", "category": "Development"},
            ],
        }
        self.customer = {"name": "Alex Morgan", "email": "alex@example.com", "organisation": "Example Ltd", "phone": "0123456789"}

    def test_all_templates_open_and_receive_known_identity_fields(self):
        for template_id, template in TEMPLATES.items():
            with self.subTest(template=template_id):
                document = Document(TEMPLATE_DIRECTORY / template["filename"])
                _populate_document(document, template_id, self.project, self.customer)
                text = self._table_text(document)
                self.assertIn("Alex Morgan", text)
                if template_id != "terms":
                    self.assertIn("Example Website", text)

    def test_project_services_populate_without_asserting_handover_completion(self):
        document = Document(TEMPLATE_DIRECTORY / TEMPLATES["handover"]["filename"])
        _populate_document(document, "handover", self.project, self.customer)
        deliverables = document.tables[4]
        self.assertEqual(deliverables.rows[1].cells[1].text, "Website design")
        self.assertEqual(deliverables.rows[1].cells[2].text, "")
        self.assertNotIn("6/28/2026", self._table_text(document))

    def test_signatures_and_signature_dates_remain_blank(self):
        document = Document(TEMPLATE_DIRECTORY / TEMPLATES["terms"]["filename"])
        _populate_document(document, "terms", self.project, self.customer)
        for table in document.tables:
            for row in table.rows:
                if row.cells[0].text.strip().casefold() in {"signature", "signature:", "date", "date:"} and len(row.cells) > 1:
                    self.assertEqual(row.cells[1].text, "")

    @staticmethod
    def _table_text(document: Document) -> str:
        return "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)


if __name__ == "__main__":
    unittest.main()
